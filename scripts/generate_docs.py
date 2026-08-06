# -*- coding: utf-8 -*-
"""Generate sprint Word documents. Run: py scripts/generate_docs.py"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"
SOPS = DOCS / "sops"
AUTHOR = "Dominic R. Nottage - KnightRoad Veritas AI Platforms"


def style_doc(doc, title, subtitle=""):
    doc.core_properties.title = title
    doc.core_properties.author = AUTHOR
    doc.add_heading(title, 0)
    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(11)
    doc.add_paragraph()


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def table_rows(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    return table


def build_sprint_doc():
    doc = Document()
    style_doc(
        doc,
        "Climate & Crisis Ops Command - 21-Day Sprint Plan",
        "Future Caribbean Buildathon 2026 | Climate Risk & Disaster Coordination",
    )

    doc.add_heading("North Star", level=1)
    doc.add_paragraph(
        "Deliver one end-to-end post-storm coordination workflow a judge can watch in five minutes: "
        "signals escalate, agents orchestrate across operator and hospital-adjacent roles, "
        "two supervisors approve via HITL, and operational outputs ship on a thin map view "
        "(alerts, checklist, comms draft, audit log)."
    )
    doc.add_paragraph(
        "Positioning: crisis coordination for Caribbean operators AND the agencies that depend "
        "on them after weather events - not a 911/CAD replacement. NEMT is the live demo "
        "vertical; architecture extends to multi-agency post-storm coordination."
    )
    doc.add_paragraph(
        "Not building: full emergency GIS/CAD or 911 dispatch replacement; Global Monitor as "
        "research-only PDFs; all seven Veritas verticals; frontier-scale model training."
    )

    doc.add_heading("Founder Credibility (Pitch)", level=1)
    doc.add_paragraph(
        "Dominic Nottage - Broward County IT: part of team that built GIS tooling helping "
        "firefighters, police, and hospital ambulances coordinate response after major weather "
        "events. This sprint productizes that class of coordination problem as agentic ops "
        "for Caribbean operators - with human approval on every high-stakes action."
    )

    doc.add_heading("MVP Definition", level=1)
    doc.add_paragraph(
        "Demo scenario: Tropical storm watch escalates in Nassau. NEMT operator runs scheduled "
        "medical transport; hospital transport liaison reviews the same incident on a shared "
        "command surface with map pins and corridor status."
    )
    numbered(
        doc,
        [
            "Signal ingest - weather + institutional threshold crossed (public feeds or labeled demo mode).",
            "Thin map layer - affected area, trip/facility pins, 1-2 closed corridors (JSON/static layers, not live ESRI/CAD).",
            "Monitor agent - brief with severity, affected geography, confidence, citations.",
            "Triage agent - ranked impacts (trips at risk, hospital partner effects, corridor conflicts).",
            "Action agent - checklist, driver alerts, hospital bulletin draft (COMMS-03).",
            "Dual HITL - NEMT dispatch supervisor AND hospital transport liaison approve or edit; nothing auto-executes.",
            "Audit log - signal to reasoning to dual approval to outputs.",
        ],
    )

    doc.add_heading("One-Line Pitch", level=2)
    doc.add_paragraph(
        "Post-storm crisis coordination for Caribbean operators and their hospital partners - "
        "agents turn institutional signals into approved operational actions on a shared command "
        "map, not another dashboard or 911 CAD clone."
    )

    for week_title, goal, rows, exit_crit in [
        (
            "Week 1 - Architecture + Skeleton (Days 1-7)",
            "Runnable command surface with sample data and Monitor agent working.",
            [
                ("Day 1", "Repo structure, env, track switch note to organizers"),
                ("Day 2", "OpenClaw spike - hello agent, tool call, logging"),
                ("Day 3", "Signal ingest module (weather + institutional or demo mock)"),
                ("Day 4", "Command UI shell - alert banner, agent timeline, dual approval panel, map panel placeholder"),
                ("Day 5", "Monitor agent: threshold to brief with citations"),
                ("Day 6", "Sample NEMT dispatch + geo layers (pins, corridors) + crisis SOP RAG"),
                ("Day 7", "Logbook #1 + architecture diagram; demo signal to Monitor brief on map"),
            ],
            "Exit criteria: Signal in, Monitor brief on screen, audit log entry.",
        ),
        (
            "Week 2 - Multi-Agent + HITL + Map (Days 8-14)",
            "Full workflow with dual human approval and thin GIS coordination view.",
            [
                ("Day 8", "Triage agent: rank impacted trips, facilities, corridor conflicts"),
                ("Day 9", "Action agent: checklist + hospital bulletin + driver comms draft"),
                ("Day 10", "Dual HITL UI - NEMT supervisor + hospital liaison roles"),
                ("Day 11", "Thin map layer - pins, zones, closed corridors synced to triage output"),
                ("Day 12", "Wire Monitor to Triage to Action; dual approval gate"),
                ("Day 13", "Audit log: steps, citations, both approvers, timestamps"),
                ("Day 14", "Polish multi-agency scenario copy; Logbook #2; 2-min screen capture"),
            ],
            "Exit criteria: End-to-end workflow with map view and dual-approved action pack.",
        ),
        (
            "Week 3 - Eval, Demo, Pitch (Days 15-21)",
            "Judge-ready demo and rubric-aligned narrative.",
            [
                ("Day 15", "Eval harness - 5-10 scripted scenarios"),
                ("Day 16", "Efficiency narrative - token/latency logging"),
                ("Day 17", "Demo script rehearsal (5 min)"),
                ("Day 18", "Defensibility + Broward credibility slide; Phase 2 CAD/EMS integration roadmap"),
                ("Day 19", "Staging deploy + backup offline demo video"),
                ("Day 20", "Logbook #3 + mentor questions list"),
                ("Day 21", "Demo day - live run + Q&A prep"),
            ],
            "Exit criteria: Repeatable 5-min demo, eval logged, pitch deck <= 8 slides.",
        ),
    ]:
        doc.add_heading(week_title, level=1)
        doc.add_paragraph(f"Goal: {goal}")
        table_rows(doc, ["Day", "Deliverable"], rows)
        doc.add_paragraph(exit_crit)

    doc.add_heading("5-Minute Demo Script", level=1)
    table_rows(
        doc,
        ["Time", "Beat", "Talk Track"],
        [
            ("0:00-0:30", "Hook", "After a storm, coordination fails across NEMT, hospitals, and corridors - not because of missing 911, but because operators cannot agree on who moves."),
            ("0:30-1:00", "Context + cred", "Storm watch escalates; I helped build GIS coordination for fire, police, and EMS after weather events at Broward County IT - this is that class of problem, agentic."),
            ("1:00-2:00", "Monitor + map", "Signal + brief with citations; map shows affected pins and a closed corridor."),
            ("2:00-3:00", "Triage + Action", "P1 dialysis trips ranked; checklist + hospital COMMS-03 draft."),
            ("3:00-4:00", "Dual HITL", "NEMT supervisor and hospital liaison each review, edit, approve; audit log captures both."),
            ("4:00-4:45", "Why you", "Veritas core; Bahamas-rooted; same command layer extends post-sprint - not rebuilding county CAD in 21 days."),
            ("4:45-5:00", "Close", "Climate coordination + operator proof + agentic multi-stakeholder HITL."),
        ],
    )

    doc.add_heading("Pitch Deck Outline (8 Slides Max)", level=1)
    numbered(
        doc,
        [
            "Title - Climate & Crisis Ops Command | Veritas | Future Caribbean",
            "Problem - post-storm fragmentation across NEMT, hospitals, corridors (Caribbean + diaspora)",
            "Founder proof - Broward County IT GIS for multi-agency weather response (fire, police, EMS/hospital)",
            "Solution - agentic command surface: signals to agents to dual HITL to approved actions on a map",
            "Demo - one workflow (NEMT + hospital liaison), not 911 CAD replacement",
            "Agentic architecture - Monitor / Triage / Action + audit + eval",
            "Defensibility - operator SOPs, workflow moat, sovereign/on-prem path (OWC prize angle)",
            "Ask / roadmap - pilot operators; Phase 2 EMS/fire/police integrations",
        ],
    )

    doc.add_heading("Logbook Entry #2 (Paste-Ready)", level=1)
    doc.add_paragraph(
        "Week 2 - Climate & Crisis Ops Command\n\n"
        "Track: Climate Risk & Disaster Coordination.\n\n"
        "Week 2 delivered: Triage agent + map sync, Action agent + COMMS-03, triple HITL "
        "(NEMT + PMH + Doctor's Hospital), orchestrator pipeline, full audit trail, "
        "multi-agency scenario polish, Logbook #2, 2-min screen capture script.\n\n"
        "Exit criteria met: end-to-end workflow with triage-synced map and triple-approved action pack.\n\n"
        "Next: Week 3 eval harness, efficiency narrative, 5-min demo rehearsal, staging deploy."
    )

    doc.add_heading("Logbook Entry #1 (Paste-Ready)", level=1)
    doc.add_paragraph(
        "Week 1 - Climate & Crisis Ops Command\n\n"
        "Track: Climate Risk & Disaster Coordination (pivot from Open Track).\n\n"
        "Sprint goal: post-storm multi-agency coordination - signals to multi-agent orchestration "
        "to dual HITL approvals to outputs on a thin map view. Demo vertical: NEMT + hospital "
        "transport liaison (not 911/CAD replacement).\n\n"
        "Founder context: Broward County IT GIS experience coordinating fire, police, and "
        "hospital ambulances after major weather events.\n\n"
        "This week: repo + architecture; OpenClaw spike; Monitor agent; command UI with map "
        "placeholder; synthetic dispatch + geo layers + crisis SOP RAG.\n\n"
        "Ask mentors: dual-approver HITL patterns; lightweight map layers for demo credibility."
    )

    doc.add_heading("Track Switch Note (Organizers)", level=1)
    doc.add_paragraph(
        "Accepted on Open Track with KnightRoad Veritas. Sprint build: Climate & Crisis Ops "
        "Command under Climate Risk & Disaster Coordination - post-storm coordination for "
        "Caribbean operators and hospital-adjacent partners. Reframes Global Monitor as "
        "operational command, not research briefs. Live demo: NEMT workflow + hospital liaison "
        "on shared map; architecture extends toward multi-agency response without claiming "
        "911/CAD replacement in the sprint window."
    )

    doc.add_heading("Scope Guards", level=1)
    bullets(
        doc,
        [
            "One workflow: storm escalation to NEMT + hospital liaison coordination on shared incident.",
            "Thin map only - pins, zones, static corridors; no full GIS/CAD or live ESRI stack.",
            "Do NOT claim 911 dispatch replacement or county-grade CAD in the sprint demo.",
            "Demo mode OK - label synthetic data clearly.",
            "Dual HITL is non-negotiable (NEMT supervisor + hospital liaison).",
            "Logbook every week.",
            "Stay on Veritas - do not join consumer tourism apps as teammate.",
        ],
    )

    doc.add_heading("Rubric Mapping", level=1)
    table_rows(
        doc,
        ["Area", "Gap Today", "Sprint Fix"],
        [
            ("Agentic AI (50%)", "Single-agent RAG", "3-agent orchestration + HITL + eval"),
            ("PMF", "Limited revenue proof", "Multi-stakeholder coordination story + design partner"),
            ("Defensibility", "Templates copyable", "SOP RAG + dual-approver workflow + Broward lineage"),
            ("Team", "Solo", "Recruit AI/ML engineer or document solo + advisory"),
            ("Efficiency", "High compute claim", "Log tokens/latency; small corpus RAG"),
        ],
    )

    doc.save(ROOT / "SPRINT.docx")


def build_architecture_doc():
    doc = Document()
    style_doc(doc, "Climate & Crisis Ops Command - System Architecture", "Multi-agent crisis ops | HITL | Audit-first")

    doc.add_heading("Overview", level=1)
    doc.add_paragraph(
        "Climate & Crisis Ops Command extends KnightRoad Veritas into post-storm multi-agency "
        "coordination. External signals and operator SOPs feed Monitor, Triage, and Action "
        "agents. A thin map layer shows geographic impact. Two human roles (NEMT dispatch "
        "supervisor and hospital transport liaison) must approve before outputs release."
    )
    doc.add_paragraph(
        "Design inspiration: county-scale GIS coordination for fire, police, and hospital EMS "
        "after major weather events (Broward County IT). Sprint scope: coordination command "
        "surface, NOT full 911/CAD replacement."
    )

    doc.add_heading("Explicitly Out of Scope (Sprint)", level=1)
    bullets(
        doc,
        [
            "911 call intake or CAD dispatch replacement",
            "Live ESRI/ArcGIS enterprise stack or AVL integrations",
            "Fire/police module builds (Phase 2 roadmap only)",
            "Real PHI in public demos",
        ],
    )

    doc.add_heading("Layers", level=1)
    for title, body in [
        (
            "Signal Layer",
            "Weather/hurricane feeds (NHC or demo mock), institutional signals (UN, World Bank, "
            "Uppsala or demo corpus), operator SOPs (RAG corpus in docs/sops/).",
        ),
        (
            "Thin GIS / Map Layer",
            "Static or JSON-driven map: service area boundary, trip/facility pins, flood-risk "
            "zones, closed corridor polylines. Purpose: coordination context for judges and "
            "multi-stakeholder HITL - not production routing engine.",
        ),
        (
            "Agent Orchestration",
            "Monitor (what changed?), Triage (what matters?), Action (what to do?). OpenClaw or "
            "equivalent with JSON schemas between agents.",
        ),
        (
            "Human-in-the-Loop (Dual Role)",
            "Role A: NEMT dispatch supervisor. Role B: hospital transport liaison. Both review "
            "same incident; both must approve (or configured sequential approve) before comms/checklist release.",
        ),
        (
            "Workflow Outputs",
            "Alerts, dispatch checklist, hospital bulletin (COMMS-03), driver comms, PDF action pack.",
        ),
        (
            "Audit and Eval",
            "Append-only log with citations and both approver IDs. Eval harness for scripted storm scenarios.",
        ),
    ]:
        doc.add_heading(title, level=2)
        doc.add_paragraph(body)

    doc.add_heading("Agent Responsibilities", level=1)
    for title, body in [
        (
            "Monitor Agent",
            "Inputs: signals, thresholds, SOP RAG. Output: situation brief with citations. "
            "Tools: fetch_signals, query_sop, write_brief.",
        ),
        (
            "Triage Agent",
            "Inputs: Monitor brief, dispatch data, geo layers (corridor closures). Output: ranked "
            "impacts across trips, facilities, and routes. Tools: load_dispatch, load_geo, "
            "score_impact, rank_items.",
        ),
        (
            "Action Agent",
            "Inputs: triage output, SOP templates. Output: checklist, alerts, comms draft. "
            "Tools: generate_checklist, draft_comms, create_tasks.",
        ),
    ]:
        doc.add_heading(title, level=2)
        doc.add_paragraph(body)

    doc.add_heading("Data Flow", level=1)
    numbered(
        doc,
        [
            "Signal threshold crossed - event logged; map highlights affected zone.",
            "Monitor Agent produces brief with geographic summary.",
            "Triage Agent ranks impacts using dispatch + corridor/facility geo data.",
            "Action Agent proposes action pack (checklist + hospital + driver comms).",
            "NEMT supervisor reviews in HITL panel - approve, reject, or edit.",
            "Hospital liaison reviews same incident - second approval required.",
            "On dual approve - outputs released; audit log finalized with both approver IDs.",
        ],
    )

    doc.add_heading("Personas (Demo)", level=1)
    table_rows(
        doc,
        ["Role", "Organization", "HITL responsibility"],
        [
            ("Dispatch Supervisor", "NEMT operator (demo)", "Trip holds, driver comms, dispatch checklist"),
            ("Hospital Transport Liaison", "Partner hospital (demo)", "P1 exceptions, COMMS-03 bulletin to partners"),
        ],
    )

    doc.add_heading("Phase 2 Roadmap (Post-Sprint)", level=1)
    bullets(
        doc,
        [
            "EMS-adjacent transport desk integrations",
            "Shelter and fleet logistics personas",
            "Fire/police situational awareness feeds (read-only)",
            "Deeper GIS routing when pilot agency available",
            "Optional on-prem sovereign stack (OWC deployment angle)",
        ],
    )

    doc.add_heading("Target Repository Layout", level=1)
    bullets(
        doc,
        [
            "src/signals/ - ingest adapters",
            "src/agents/ - monitor, triage, action",
            "src/orchestrator/ - pipeline runner",
            "src/hitl/ - dual-role approval API",
            "src/audit/ - log writer",
            "src/geo/ - map layers, pins, corridors (thin GIS)",
            "src/ui/ - command surface with map panel",
            "data/sample-dispatch.csv - synthetic trips",
            "data/geo/ - corridors, facilities, zones (JSON/GeoJSON)",
            "docs/sops/ - operator SOPs for RAG",
        ],
    )

    doc.add_heading("Tech Stack (Sprint)", level=1)
    bullets(
        doc,
        [
            "Node.js / JavaScript - API + UI (matches Veritas parent)",
            "OpenClaw - multi-agent orchestration",
            "LLM APIs - tool calling with token logging",
            "RAG - embeddings over SOP documents",
            "Optional H200 - eval or open-model inference if keys assigned",
        ],
    )

    doc.add_heading("Security and Demo Mode", level=1)
    doc.add_paragraph(
        "Production may include PHI/PII - treat pilot data as confidential. Buildathon demo uses "
        "synthetic dispatch data with DEMO MODE labeled in UI. HITL applies in both modes."
    )

    doc.add_heading("Parent Platform", level=1)
    doc.add_paragraph(
        "KnightRoad Veritas - knightroadveritas.app. Reuse NEMT catalog, Global Monitor signal "
        "concept, command-center UI patterns."
    )

    doc.save(DOCS / "architecture.docx")


def build_sop_template():
    doc = Document()
    style_doc(
        doc,
        "NEMT and Multi-Agency Crisis Coordination - SOP Template",
        "NEMT operator + hospital partner roles | Replace bracketed fields before production use",
    )

    table_rows(
        doc,
        ["Field", "Value"],
        [
            ("Organization", "[Operator Name]"),
            ("Depot / Region", "[Region]"),
            ("SOP ID", "NEMT-CRISIS-001"),
            ("Version", "1.0"),
            ("Effective Date", "[YYYY-MM-DD]"),
            ("Owner", "[Operations Manager]"),
        ],
    )

    doc.add_heading("1. Purpose", level=1)
    doc.add_paragraph(
        "Define coordinated response during weather watches, storm warnings, and recovery between "
        "[NEMT Operator] and partner hospitals/EOC liaisons. Ensures patient and driver safety, "
        "aligned comms to hospital partners, and compliance with care agreements. Does not "
        "replace 911 or municipal CAD."
    )

    doc.add_heading("2. Scope", level=1)
    bullets(
        doc,
        [
            "Scheduled NEMT trips during declared weather events.",
            "NEMT dispatch, drivers, supervisors; hospital transport liaisons.",
            "Shared situational awareness via map/corridor status (when command system deployed).",
            "Excludes emergency 911 transport and fire/police dispatch.",
        ],
    )

    doc.add_heading("3. Escalation Thresholds", level=1)
    bullets(
        doc,
        [
            "Level 1 Monitor: [e.g. Tropical Storm Watch for service area]",
            "Level 2 Prepare: [e.g. wind forecast > 40 mph]",
            "Level 3 Restrict: [e.g. Hurricane Warning or road closures]",
            "Level 4 Suspend: [e.g. curfew or sustained winds > 55 mph]",
        ],
    )

    doc.add_heading("4. Roles", level=1)
    bullets(
        doc,
        [
            "NEMT Dispatch Supervisor: approve trip holds, driver comms, dispatch checklist.",
            "Hospital Transport Liaison: approve P1 exceptions and hospital bulletin (COMMS-03).",
            "Dispatcher: execute approved checklist after dual approval where required.",
            "Driver: report conditions; follow hold directives.",
            "Clinical Liaison: prioritize time-sensitive appointments.",
        ],
    )

    doc.add_heading("5. Map / Corridor Coordination", level=1)
    bullets(
        doc,
        [
            "Maintain list of flood-prone corridors and bridge-dependent routes.",
            "When corridor marked CLOSED in command system, no P3/P4 trips routed through it.",
            "Supervisor and liaison review map pins before approving action pack.",
        ],
    )

    doc.add_heading("6. Level 2 Prepare Actions", level=1)
    numbered(
        doc,
        [
            "Pull 24-hour trip manifest; flag flood-zone routes.",
            "Pre-notify patients; confirm callback numbers.",
            "Stage driver comms; verify vehicle readiness.",
        ],
    )

    doc.add_heading("7. Level 3 Restrict Actions", level=1)
    numbered(
        doc,
        [
            "Hold non-essential trips; liaison approves exceptions.",
            "Consolidate remaining trips where safe.",
            "Require driver check-in every [X] hours.",
        ],
    )

    doc.add_heading("8. Communication Templates", level=1)
    doc.add_paragraph("Patient hold (template):")
    doc.add_paragraph(
        '"[Operator] update: Due to [weather event], your trip on [date/time] is on hold. '
        'We will contact you at [phone]. For urgent medical needs, contact your provider or 911."'
    )

    doc.add_heading("9. Agent Integration", level=1)
    doc.add_paragraph(
        "Climate & Crisis Ops Command agents may propose checklists and comms from this SOP. "
        "NEMT supervisor and hospital liaison must both approve before notifications are sent. "
        "Agents cite SOP section IDs in the audit log. System does not replace 911 or CAD."
    )

    doc.save(SOPS / "nemt-crisis-sop-template.docx")


def build_sample_sop():
    doc = Document()
    style_doc(
        doc,
        "Nassau Metro NEMT - Multi-Agency Storm Coordination (Sample)",
        "NEMT + hospital liaison demo | DEMO DATA ONLY | Not 911/CAD",
    )

    doc.add_paragraph(
        "Sample SOP for sprint demo and RAG testing. Covers NEMT operator and partner hospital "
        "transport liaison on a shared incident. Not a real operator policy. Use synthetic "
        "dispatch and geo data in public demos."
    )

    table_rows(
        doc,
        ["Field", "Value"],
        [
            ("Organization", "Nassau Metro Medical Transport (Demo)"),
            ("Hospital Partner", "Princess Margaret Hospital Transport Desk (Demo)"),
            ("Region", "New Providence, Bahamas"),
            ("SOP ID", "NEMT-CRISIS-001-DEMO"),
            ("NEMT Owner", "Maria Clarke, Operations Supervisor"),
            ("Hospital Liaison", "James Rolle, Transport Coordinator (Demo)"),
        ],
    )

    doc.add_heading("Demo Map Corridors (Geo Layer)", level=1)
    bullets(
        doc,
        [
            "CORR-01 Paradise Island Bridge approach - flood risk; CLOSED at Level 3.",
            "CORR-02 Eastern Road low segment - flood risk; RESTRICTED at Level 2+.",
            "FAC-01 Princess Margaret Hospital - partner facility pin.",
            "FAC-02 Nassau Metro Depot - operator base pin.",
        ],
    )

    doc.add_heading("Dual Approval Rule (Level 3+)", level=1)
    doc.add_paragraph(
        "P1 trip proceeding during Level 3 Restrict requires approval from BOTH NEMT Dispatch "
        "Supervisor (Maria Clarke) AND Hospital Transport Liaison (James Rolle). COMMS-03 "
        "hospital bulletin requires liaison approval before send."
    )

    doc.add_heading("Escalation Thresholds (Demo)", level=1)
    bullets(
        doc,
        [
            "Level 1 Monitor: NHC Tropical Storm Watch for New Providence OR winds > 35 mph in 72h.",
            "Level 2 Prepare: Tropical Storm Warning OR > 40 trips with flood-prone routes.",
            "Level 3 Restrict: Hurricane Watch OR bridge/traffic advisory on primary corridors.",
            "Level 4 Suspend: Hurricane Warning, curfew, or winds > 55 mph at depot.",
        ],
    )

    doc.add_heading("Level 2 Prepare Actions", level=1)
    numbered(
        doc,
        [
            "Export 24h manifest; tag dialysis/oncology as PRIORITY.",
            "Flag 8 flood-risk stops (Paradise Island approach, low-lying routes).",
            "SMS pre-notify patients in first 12h window.",
            "Brief drivers; confirm fuel > 75% and emergency kits.",
        ],
    )

    doc.add_heading("Level 3 Restrict Actions", level=1)
    numbered(
        doc,
        [
            "Hold non-PRIORITY trips until clinical liaison approval.",
            "Consolidate approved trips; no overnight single-driver routes.",
            "Driver check-in every 2 hours.",
            "Use COMMS-03 for hospital partners; HITL required before send.",
        ],
    )

    doc.add_heading("Priority Trip Codes", level=1)
    table_rows(
        doc,
        ["Code", "Meaning", "Level 3 Rule"],
        [
            ("P1", "Dialysis / life-sustaining", "Proceed with supervisor + liaison approval"),
            ("P2", "Oncology / time-sensitive", "Hold unless 6h window critical"),
            ("P3", "Routine appointment", "Hold automatically at Level 3"),
            ("P4", "Discharge / return", "Reschedule unless hospital escalation"),
        ],
    )

    doc.add_heading("COMMS-03 Hospital Partner Bulletin", level=2)
    doc.add_paragraph(
        "Subject: Nassau Metro NEMT - Operations Level 3 Restrict\n\n"
        "We entered Level 3 due to [event]. Non-priority trips on hold. P1 trips need dual "
        "approval. Route urgent requests through dispatch. Recovery window TBD post all-clear."
    )

    doc.add_heading("COMMS-04 Patient Trip Hold", level=2)
    doc.add_paragraph(
        "Nassau Metro Transport: Your trip on [DATE/TIME] is on hold due to storm conditions. "
        "We will call [PHONE] to reschedule. For emergency, dial 911."
    )

    doc.add_heading("Agent Citation IDs for RAG", level=1)
    bullets(
        doc,
        [
            "NEMT-CRISIS-001-DEMO Thresholds",
            "NEMT-CRISIS-001-DEMO Map Corridors CORR-01 CORR-02",
            "NEMT-CRISIS-001-DEMO Dual Approval Rule",
            "NEMT-CRISIS-001-DEMO Level 2 Actions",
            "NEMT-CRISIS-001-DEMO Level 3 Actions",
            "NEMT-CRISIS-001-DEMO Priority Codes P1-P4",
            "NEMT-CRISIS-001-DEMO COMMS-03",
            "NEMT-CRISIS-001-DEMO COMMS-04",
        ],
    )

    doc.save(SOPS / "nassau-metro-storm-escalation-sample.docx")


def build_track_switch_doc():
    doc = Document()
    style_doc(
        doc,
        "Track Switch Note - Future Caribbean Organizers",
        "Paste into portal message, WhatsApp, or email to program staff",
    )

    doc.add_paragraph("To: Future Caribbean Buildathon organizers")
    doc.add_paragraph("From: Dominic R. Nottage | KnightRoad Veritas AI Platforms")
    doc.add_paragraph("Re: Sprint track alignment - Climate & Crisis Ops Command")
    doc.add_paragraph("")

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(
        "I was accepted on Open Track with KnightRoad Veritas. For the 21-day sprint I am "
        "building Climate & Crisis Ops Command under Climate Risk & Disaster Coordination."
    )

    doc.add_heading("What Changed", level=1)
    bullets(
        doc,
        [
            "Same Veritas agent + command-center architecture and founder team (solo).",
            "Reframes Global Monitor from institutional research briefs to operational crisis command.",
            "Live demo: NEMT operator + hospital transport liaison on a shared map with dual HITL.",
            "Does NOT claim 911/CAD replacement in the sprint window.",
        ],
    )

    doc.add_heading("Why Climate Track", level=1)
    doc.add_paragraph(
        "Aligns with program climate priorities and founder experience: GIS coordination for "
        "fire, police, and hospital EMS after major weather events (Broward County IT). "
        "Product is post-storm multi-agency coordination for Caribbean operators - not a "
        "consumer tourism app."
    )

    doc.add_heading("Sprint MVP (Demo Day)", level=1)
    doc.add_paragraph(
        "One workflow: storm signal escalates, Monitor/Triage/Action agents propose an action "
        "pack, NEMT supervisor and hospital liaison both approve, audit log captures citations "
        "and approvals. Thin map layer shows corridors and facility pins (demo/synthetic data)."
    )

    doc.add_heading("Contact", level=1)
    doc.add_paragraph(
        "Dominic R. Nottage | dnottage.veritasplatform@gmail.com | knightroadveritas.app | Nassau, Bahamas"
    )

    doc.save(DOCS / "track-switch-note.docx")


def build_openclaw_spike_doc():
    doc = Document()
    style_doc(
        doc,
        "Day 2 - OpenClaw Spike Notes",
        "Future Caribbean sprint | Monitor agent | Tool loop + logging",
    )

    doc.add_heading("Decision", level=1)
    doc.add_paragraph(
        "Use an OpenClaw-compatible tool-calling agent loop inside Express for Week 1. "
        "Install OpenClaw Gateway globally when compute/mentor keys arrive; skill stub lives "
        "at openclaw/workspace/skills/climate-monitor/."
    )

    doc.add_heading("What Was Built", level=1)
    numbered(
        doc,
        [
            "src/agents/runtime/tools.js - get_signal_status, query_sop, summarize_dispatch",
            "src/agents/runtime/logger.js - structured agent event log",
            "src/agents/runtime/runAgent.js - Monitor spike loop (demo or OpenAI)",
            "POST /api/agents/monitor/spike - HTTP entry point",
            "GET /api/agents/logs - audit-style event stream",
            "npm run spike:monitor - CLI test",
        ],
    )

    doc.add_heading("OpenClaw Gateway (Optional Next)", level=1)
    numbered(
        doc,
        [
            "npm install -g openclaw@latest",
            "openclaw onboard --install-daemon",
            "Symlink skill to ~/.openclaw/workspace/skills/climate-monitor",
            "openclaw doctor && openclaw gateway status",
        ],
    )

    doc.add_heading("Week 2 Wiring", level=1)
    bullets(
        doc,
        [
            "Add Triage and Action agents to same runtime pattern",
            "Connect signal ingest to get_signal_status (replace demo mock)",
            "Dual HITL gate before action pack release",
            "Optional: delegate agent turns to OpenClaw Gateway",
        ],
    )

    doc.save(DOCS / "openclaw-spike.docx")


def main():
    DOCS.mkdir(parents=True, exist_ok=True)
    SOPS.mkdir(parents=True, exist_ok=True)
    build_sprint_doc()
    build_architecture_doc()
    build_sop_template()
    build_sample_sop()
    build_track_switch_doc()
    build_openclaw_spike_doc()
    print("Created:")
    for p in [
        ROOT / "SPRINT.docx",
        DOCS / "architecture.docx",
        DOCS / "track-switch-note.docx",
        DOCS / "openclaw-spike.docx",
        SOPS / "nemt-crisis-sop-template.docx",
        SOPS / "nassau-metro-storm-escalation-sample.docx",
    ]:
        print(f"  {p}")


if __name__ == "__main__":
    main()
